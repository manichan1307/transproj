# main.py - Final Kivy Application with Offline Translation
import os
import threading
from datetime import datetime
import requests  # For downloading models

# Kivy UI components
from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.label import Label
from kivy.uix.button import Button
from kivy.uix.spinner import Spinner
from kivy.uix.progressbar import ProgressBar
from kivy.uix.scrollview import ScrollView
from kivy.clock import Clock
from kivy.utils import platform

# Core libraries for document manipulation
from docx import Document

# Core NLP/ML libraries for offline translation
import torch
from transformers import MarianMTModel, MarianTokenizer

# --- Global Configuration & State ---
APP_NAME = "DocTranslator"
MODEL_BASE_URL = "https://huggingface.co/Helsinki-NLP/"
MODELS_TO_DOWNLOAD = {
    "opus-mt-it-en": ["config.json", "pytorch_model.bin", "source.spm", "target.spm", "tokenizer_config.json",
                      "vocab.json"],
    "opus-mt-fr-en": ["config.json", "pytorch_model.bin", "source.spm", "target.spm", "tokenizer_config.json",
                      "vocab.json"],
    "opus-mt-en-fr": ["config.json", "pytorch_model.bin", "source.spm", "target.spm", "tokenizer_config.json",
                      "vocab.json"],
}

# Global NMT model, tokenizer, and device for efficiency
_nmt_model = None
_nmt_tokenizer = None
_nmt_device = None
_last_model_path = None


def get_storage_path():
    """Get the primary writable storage path."""
    if platform == 'android':
        from android.storage import primary_external_storage_path
        return primary_external_storage_path()
    return os.path.expanduser('~')


def get_app_dir():
    """Get the application's private data directory."""
    if platform == 'android':
        from kivy.utils import get_user_data_dir
        return get_user_data_dir(APP_NAME)
    return os.path.join(os.path.expanduser('~'), f'.{APP_NAME}')


def get_model_path(model_name):
    """Get the full path to a specific model's directory."""
    return os.path.join(get_app_dir(), "offline_model", model_name)


def load_nmt_model(model_path):
    """Loads the NMT model and tokenizer from a local path."""
    global _nmt_model, _nmt_tokenizer, _nmt_device, _last_model_path
    if _nmt_model is None or _last_model_path != model_path:
        try:
            _nmt_tokenizer = MarianTokenizer.from_pretrained(model_path)
            _nmt_model = MarianMTModel.from_pretrained(model_path)
            _nmt_device = torch.device("cpu")  # Force CPU for mobile compatibility
            _nmt_model.to(_nmt_device)
            _last_model_path = model_path
            print(f"[INFO] Loaded model: {model_path}")
        except Exception as e:
            print(f"[ERROR] Failed to load model from {model_path}. Error: {e}")
            raise e
    return _nmt_model, _nmt_tokenizer, _nmt_device


def translate_batch_text(texts, tokenizer, model, device):
    """Translates a list of texts in batches for efficiency."""
    texts = [str(t) for t in texts if t and str(t).strip()]
    if not texts:
        return []
    translated_texts = []
    with torch.no_grad():  # Inference mode is crucial
        for i in range(0, len(texts), 16):  # Process in batches of 16
            batch_texts = texts[i:i + 16]
            encoded = tokenizer(batch_texts, return_tensors="pt", padding=True, truncation=True, max_length=512).to(
                device)
            outputs = model.generate(**encoded)
            translated_batch = [tokenizer.decode(g, skip_special_tokens=True) for g in outputs]
            translated_texts.extend(translated_batch)
    return translated_texts


def process_docx_text_only(input_path, output_folder, model, tokenizer, device, src_lang, tgt_lang, log_callback):
    """Translates a DOCX document, focusing ONLY on text content."""
    log_callback(f"Processing DOCX: {os.path.basename(input_path)}")
    doc = Document(input_path)
    out_doc = Document()

    texts_to_translate = []
    element_map = []

    log_callback("Pass 1: Extracting text from document...")
    for para in doc.paragraphs:
        if para.text.strip():
            texts_to_translate.append(para.text)
            element_map.append({'type': 'paragraph', 'obj': para, 'text_idx': len(texts_to_translate) - 1})
        else:
            element_map.append({'type': 'empty_paragraph'})

    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                if cell.text.strip():
                    texts_to_translate.append(cell.text)
                    row_data.append({'text': cell.text, 'text_idx': len(texts_to_translate) - 1})
                else:
                    row_data.append({'text': ''})
            table_data.append(row_data)
        element_map.append({'type': 'table', 'obj': table, 'data': table_data})

    log_callback(f"Pass 2: Translating {len(texts_to_translate)} text segments...")
    translated_texts = translate_batch_text(texts_to_translate, tokenizer, model, device)
    log_callback("Translation complete. Reconstructing document...")

    for element in element_map:
        if element['type'] == 'paragraph':
            p = out_doc.add_paragraph()
            translated_text = translated_texts[element['text_idx']]
            run = p.add_run(translated_text)
            orig_run = element['obj'].runs[0] if element['obj'].runs else None
            if orig_run:
                run.bold = orig_run.bold
                run.italic = orig_run.italic
                run.underline = orig_run.underline

        elif element['type'] == 'empty_paragraph':
            out_doc.add_paragraph()

        elif element['type'] == 'table':
            orig_table = element['obj']
            new_table = out_doc.add_table(rows=len(orig_table.rows), cols=len(orig_table.columns),
                                          style=orig_table.style)
            for r_idx, row_data in enumerate(element['data']):
                for c_idx, cell_data in enumerate(row_data):
                    if 'text_idx' in cell_data:
                        new_table.cell(r_idx, c_idx).text = translated_texts[cell_data['text_idx']]

    name_no_ext = os.path.splitext(os.path.basename(input_path))[0]
    out_doc_path = os.path.join(output_folder, f"{name_no_ext}_{src_lang}_{tgt_lang}_translated.docx")
    out_doc.save(out_doc_path)
    log_callback(f"✅ Success! Saved to: {os.path.basename(out_doc_path)}")
    return out_doc_path


class TranslatorLayout(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.orientation = 'vertical'
        self.padding = [20, 20, 20, 20]
        self.spacing = 10
        self.input_file = None
        self.output_folder = os.path.join(get_storage_path(), 'TranslatedDocuments')
        os.makedirs(self.output_folder, exist_ok=True)

        self.add_widget(Label(text="Offline Document Translator", font_size='24sp', size_hint_y=None, height=40))

        self.file_label = Label(text="No file selected.", size_hint_y=None, height=40)
        self.add_widget(self.file_label)
        self.browse_btn = Button(text="Select DOCX File", on_press=self.select_file, size_hint_y=None, height=50)
        self.add_widget(self.browse_btn)

        lang_box = BoxLayout(orientation='horizontal', size_hint_y=None, height=50)
        self.from_spinner = Spinner(text='Italian', values=('Italian', 'French', 'English'), size_hint_x=0.4)
        self.to_spinner = Spinner(text='English', values=('English', 'French', 'Italian'), size_hint_x=0.4)
        swap_btn = Button(text='↔', on_press=self.swap_languages, size_hint_x=0.2)
        lang_box.add_widget(self.from_spinner)
        lang_box.add_widget(swap_btn)
        lang_box.add_widget(self.to_spinner)
        self.add_widget(lang_box)

        self.translate_btn = Button(text="Start Translation", on_press=self.start_translation_thread, size_hint_y=None,
                                    height=50)
        self.add_widget(self.translate_btn)

        self.progress_bar = ProgressBar(max=100, size_hint_y=None, height=20)
        self.add_widget(self.progress_bar)

        self.log_box = Label(text="Logs will appear here.", markup=True, halign='left', valign='top')
        self.log_box.bind(size=self.log_box.setter('text_size'))
        scroll_view = ScrollView()
        scroll_view.add_widget(self.log_box)
        self.add_widget(scroll_view)

    def select_file(self, instance):
        if platform == 'android':
            from android.permissions import request_permissions, Permission
            request_permissions(
                [Permission.READ_EXTERNAL_STORAGE, Permission.WRITE_EXTERNAL_STORAGE, Permission.INTERNET])
            self.log_message(f"INFO: For this demo, place your DOCX file in the folder: {self.output_folder}")
            test_file_path = os.path.join(self.output_folder, "test.docx")  # Example file
            if os.path.exists(test_file_path):
                self.input_file = test_file_path
                self.file_label.text = f"Using test file: test.docx"
            else:
                self.log_message(
                    f"WARN: Please place a file named 'test.docx' in the TranslatedDocuments folder on your phone.")
                self.file_label.text = "Test file not found!"
        else:  # For desktop testing
            from tkinter import Tk, filedialog
            root = Tk()
            root.withdraw()
            file_path = filedialog.askopenfilename(filetypes=[("Word Document", "*.docx")])
            if file_path:
                self.input_file = file_path
                self.file_label.text = f"Selected: {os.path.basename(file_path)}"

    def swap_languages(self, instance):
        from_lang, to_lang = self.from_spinner.text, self.to_spinner.text
        self.from_spinner.text = to_lang
        self.to_spinner.text = from_lang

    def log_message(self, msg):
        Clock.schedule_once(lambda dt: self._update_log(msg))

    def _update_log(self, msg):
        self.log_box.text += f"\n{datetime.now().strftime('%H:%M:%S')} - {msg}"

    def update_progress(self, value):
        Clock.schedule_once(lambda dt: setattr(self.progress_bar, 'value', value))

    def start_translation_thread(self, instance):
        if not self.input_file or not os.path.exists(self.input_file):
            self.log_message("[ERROR] Please select a valid input file first.")
            return

        self.translate_btn.disabled = True
        self.log_box.text = "Starting..."
        threading.Thread(target=self.run_translation).start()

    def run_translation(self):
        try:
            lang_from_ui = self.from_spinner.text.lower()
            lang_to_ui = self.to_spinner.text.lower()
            lang_map = {"english": "en", "italian": "it", "french": "fr"}

            src_lang = lang_map[lang_from_ui]
            tgt_lang = lang_map[lang_to_ui]
            model_name = f"opus-mt-{src_lang}-{tgt_lang}"

            self.log_message(f"Selected model: {model_name}")
            self.update_progress(10)

            model_path = get_model_path(model_name)
            if not os.path.exists(model_path):
                self.log_message("Model not found locally. Starting download...")
                self.download_model(model_name)

            self.log_message("Loading translation model...")
            self.update_progress(50)
            model, tokenizer, device = load_nmt_model(model_path)

            self.log_message("Starting document processing...")
            self.update_progress(70)

            process_docx_text_only(
                self.input_file, self.output_folder, model, tokenizer, device,
                src_lang, tgt_lang, self.log_message
            )

            self.log_message("\n--- Translation Finished ---")
            self.update_progress(100)

        except Exception as e:
            self.log_message(f"[FATAL ERROR] {e}")
            import traceback
            self.log_message(traceback.format_exc())
        finally:
            Clock.schedule_once(lambda dt: setattr(self.translate_btn, 'disabled', False))

    def download_model(self, model_name):
        files_to_download = MODELS_TO_DOWNLOAD.get(model_name)
        if not files_to_download:
            raise ValueError(f"Model '{model_name}' is not configured for download.")

        model_dir = get_model_path(model_name)
        os.makedirs(model_dir, exist_ok=True)

        for i, filename in enumerate(files_to_download):
            url = f"{MODEL_BASE_URL}{model_name}/resolve/main/{filename}"
            self.log_message(f"Downloading {filename}...")
            try:
                response = requests.get(url, stream=True)
                response.raise_for_status()
                with open(os.path.join(model_dir, filename), "wb") as f:
                    for chunk in response.iter_content(chunk_size=8192):
                        f.write(chunk)
                progress = 10 + ((i + 1) / len(files_to_download)) * 40
                self.update_progress(progress)
            except requests.exceptions.RequestException as e:
                self.log_message(f"Failed to download {filename}. Error: {e}")
                raise e


class DocTranslatorApp(App):
    def build(self):
        return TranslatorLayout()


if __name__ == '__main__':
    DocTranslatorApp().run()
